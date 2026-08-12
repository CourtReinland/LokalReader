"""DOCX parser (nice-to-have)."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from lokalreader.parsers.base import ParsedChapter, ParsedDocument
from lokalreader.parsers.txt import _split_chapters


def parse_docx(path: Path) -> ParsedDocument:
    doc = Document(str(path))
    title = path.stem.replace("_", " ").title()
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            paragraphs.append("")
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or "title" in style:
            paragraphs.append(f"# {text}")
            if "title" in style and title == path.stem.replace("_", " ").title():
                title = text
        else:
            paragraphs.append(text)
    full = "\n".join(paragraphs).strip()
    return ParsedDocument(title=title, chapters=_split_chapters(full, title), format="docx")
